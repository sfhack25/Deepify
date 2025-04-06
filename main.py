"""
Deep Learner backend (MongoDB Atlas edition)
-------------------------------------------------
Now updated to use **genai.Client** with the *gemini‑2.0‑flash* model for every
Google Gemini request (instead of the older `GenerativeModel` interface).

Quick start:
  1.  export GOOGLE_API_KEY="<your_gemini_key>"
  2.  export MONGODB_URI="mongodb+srv://<user>:<pass>@cluster0.xxxxx.mongodb.net/?retryWrites=true&w=majority"
  3.  pip install -r requirements.txt
  4.  uvicorn main:app --reload

Dependencies (requirements.txt):
  fastapi
  uvicorn[standard]
  motor                 # async MongoDB driver
  pydantic              # FastAPI uses it internally
  google-generativeai   # Gemini client
  pdfplumber            # PDF text extraction
  python-docx           # DOCX text extraction
"""

import os
import json
from dotenv import load_dotenv
import tempfile
from datetime import datetime, date, timedelta
from typing import List, Optional
import re
import pdfplumber
import io
import docx
# import google.generativeai as genai
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from motor.motor_asyncio import AsyncIOMotorClient, AsyncIOMotorGridFSBucket
from bson import ObjectId
from google import genai
from fastapi import Body
from enum import Enum
import gridfs
from notes_quiz_generator import create_notes_quiz_endpoint
load_dotenv()


try:
    import docx  # type: ignore
except ImportError:
    docx = None  # will raise later if user uploads DOCX without python-docx installed

# ---------------------------------------------------------------------------
# Environment & third‑party setup
# ---------------------------------------------------------------------------
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
# Use MongoDB Atlas
MONGODB_URI = os.getenv("MONGO_URI")

if not GOOGLE_API_KEY:
    raise RuntimeError("Set GOOGLE_API_KEY environment variable")

if not MONGODB_URI:
    raise RuntimeError("Set MONGO_URI environment variable")

# Create a single Gemini client instance (new style)
GENAI_CLIENT = genai.Client(api_key=GOOGLE_API_KEY)

# Initialize MongoDB variables
mongo_client = None
db = None
fs = None

# ---------------------------------------------------------------------------
# Pydantic models (for request / response bodies)
# ---------------------------------------------------------------------------
class RoadmapEntry(BaseModel):
    date: Optional[str] = None  # ISO date or week number
    topic: str
    preQuizPrompt: Optional[str] = None
    assignment: Optional[str] = None


class CourseCreateResponse(BaseModel):
    id: str = Field(..., alias="_id")
    roadmap: List[RoadmapEntry]


# ---------------------------------------------------------------------------
# FastAPI app
# ---------------------------------------------------------------------------
app = FastAPI(title="Deep Learner API – MongoDB Atlas MVP")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class RatingEnum(str, Enum):
    easy = "easy"
    medium = "medium"
    hard = "hard"
    dont_know = "dont_know"

rating_schedule = {
    "easy": 7,
    "medium": 3,
    "hard": 1,
    "dont_know": 1
}

# ---------------------------------------------------------------------------
# Startup and shutdown events
# ---------------------------------------------------------------------------
@app.on_event("startup")
async def startup_db_client():
    global mongo_client, db, fs
    
    # Connect to MongoDB Atlas
    try:
        print(f"Attempting to connect to MongoDB Atlas at: {MONGODB_URI}")
        # Disable SSL certificate verification for MongoDB Atlas
        mongo_client = AsyncIOMotorClient(MONGODB_URI, serverSelectionTimeoutMS=5000, tlsAllowInvalidCertificates=True)
        # Test the connection
        print("Testing MongoDB connection...")
        server_info = await mongo_client.server_info()
        print(f"MongoDB server info: {server_info}")
        db = mongo_client["deep_learner"]
        # Initialize AsyncIOMotorGridFSBucket instead of synchronous GridFSBucket
        fs = AsyncIOMotorGridFSBucket(db)
        print("Successfully connected to MongoDB Atlas")
        
        # Test database access by listing collections
        collections = await db.list_collection_names()
        print(f"Available collections: {collections}")
    except Exception as e:
        print(f"Failed to connect to MongoDB Atlas: {e}")
        raise RuntimeError(f"Failed to connect to MongoDB Atlas: {e}")

@app.on_event("shutdown")
async def shutdown_db_client():
    global mongo_client
    if mongo_client:
        mongo_client.close()

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def extract_text_from_file(path: str, suffix: str) -> str:
    """Return plaintext from a PDF or DOCX syllabus file."""

    if suffix == ".pdf":
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    if suffix == ".docx":
        if docx is None:
            raise RuntimeError("python-docx not installed; cannot parse DOCX files")
        document = docx.Document(path)
        return "\n".join(p.text for p in document.paragraphs)

    raise ValueError("Unsupported file type")


def generate_roadmap(syllabus_text: str):
    """Call Gemini (gemini‑2.0‑flash) to convert raw syllabus into a JSON roadmap."""

    
    prompt = (
        "You are an expert academic planner. Given the following semester syllabus, "
        "produce a valid JSON array where each element has: `date` (YYYY-MM-DD), "
        "has element `topic` and `preQuizPrompt` which is detailed 2 sentence of what the topic is about and `assignment` due on that day if theres any pure json array please with no other info starts with { and ends with }"
        f"\n\nSyllabus:\n{syllabus_text}"
        )

    try:
        print(f"Sending prompt to Gemini API: {prompt[:200]}...")
        response = GENAI_CLIENT.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt,
        )
        
        # Log the full response
        print(f"Gemini API Response: {response}")
        
        # Depending on library version, the text may be under `text` or `content`.
        raw_text = getattr(response, "text", None) or getattr(response, "content", "")
        print(f"Extracted text from response: {raw_text}")
        
        # Clean up the response by removing markdown code block formatting
        cleaned_text = raw_text
        if "```json" in cleaned_text:
            cleaned_text = cleaned_text.split("```json")[1]
        if "```" in cleaned_text:
            cleaned_text = cleaned_text.split("```")[0]
        
        # Strip any leading/trailing whitespace
        cleaned_text = cleaned_text.strip()
        print(f"Cleaned text for JSON parsing: {cleaned_text}")
        
    except Exception as err:
        print(f"Error calling Gemini API: {err}")
        raise ValueError(f"Gemini API call failed: {err}")

    try:
        roadmap_data = json.loads(cleaned_text)
        print(f"Successfully parsed JSON roadmap: {roadmap_data}")
        return roadmap_data
    except Exception as exc:
        print(f"Error parsing JSON from Gemini response: {exc}")
        print(f"Raw response that failed to parse: {raw_text}")
        raise ValueError(
            f"Gemini returned invalid JSON: {exc}\nRaw response: {raw_text}"
        )


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------
@app.post("/courses/", response_model=CourseCreateResponse, summary="Create a course and upload syllabus")
async def create_course(name: str, syllabus: UploadFile = File(...)):
    """Accept a PDF or DOCX syllabus, extract text, generate roadmap, persist in MongoDB."""

    suffix = os.path.splitext(syllabus.filename)[-1].lower()
    if suffix not in {".pdf", ".docx"}:
        raise HTTPException(status_code=400, detail="Only PDF or DOCX files are supported")

    # Save to a temporary file to allow library parsing
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await syllabus.read())
        tmp_path = tmp.name

    # 1️ Extract text
    try:
        syllabus_text = extract_text_from_file(tmp_path, suffix)
    except Exception as err:
        raise HTTPException(status_code=500, detail=f"Failed to parse syllabus: {err}")

    # 2️ Generate roadmap via Gemini
    try:
        roadmap_json = generate_roadmap(syllabus_text)
    except ValueError as err:
        raise HTTPException(status_code=500, detail=str(err))

    # 3️ Persist to MongoDB
    course_doc = {
        "name": name,
        "created_at": datetime.utcnow(),
        "roadmap": roadmap_json,
    }
    
    try:
        # Insert the document
        print(f"Inserting course document: {name}")
        result = await db.courses.insert_one(course_doc)
        course_id = str(result.inserted_id)
        print(f"Successfully inserted course with ID: {course_id}")
        
        # Verify the document was inserted
        inserted_doc = await db.courses.find_one({"_id": result.inserted_id})
        if inserted_doc:
            print(f"Verified document insertion: {inserted_doc['name']}")
        else:
            print("Warning: Could not verify document insertion")
    except Exception as e:
        print(f"Error inserting document to MongoDB: {e}")
        raise RuntimeError(f"Error inserting document to MongoDB: {e}")
    
    # Create response with proper _id field
    response_data = {
        "_id": course_id,
        "roadmap": roadmap_json
    }
    
    return response_data


@app.get(
    "/courses/{course_id}/roadmap",
    response_model=List[RoadmapEntry],
    summary="Fetch roadmap for a course by ID",
)
async def get_roadmap(course_id: str):
    try:
        obj_id = ObjectId(course_id)
        course = await db.courses.find_one({"_id": obj_id}, {"_id": 0, "roadmap": 1})
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
        return course["roadmap"]
    except Exception as e:
        if "Invalid course ID" in str(e):
            raise HTTPException(status_code=400, detail="Invalid course ID")
        raise HTTPException(status_code=500, detail=f"Error retrieving roadmap: {str(e)}")


@app.post("/courses/{course_id}/quizzes/pre", summary="Generate flashcard-style pre-lecture quiz for all roadmap topics")
async def generate_pre_quiz(course_id: str):
    try:
        obj_id = ObjectId(course_id)
        course = await db.courses.find_one({"_id": obj_id})
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")

        roadmap = course.get("roadmap", [])
        if not roadmap:
            raise HTTPException(status_code=404, detail="No roadmap entries found")

        all_quizzes = []

        for entry in roadmap:
            topic = entry.get("topic")
            prompt = entry.get("preQuizPrompt")

            if not topic or not prompt:
                continue  # skip entries without quiz input data

            quiz_prompt = f"""
Create 10 flashcard-style questions to help a student study this topic:

Topic: {topic}
Prompt: {prompt}

Each question should be a simple question or recall prompt, and each should have a clear answer.
Return ONLY valid JSON in this format:
[
  {{
    "question": "...",
    "answer": "..."
  }},
  ...
]
"""

            try:
                response = GENAI_CLIENT.models.generate_content(
                    model="gemini-2.0-flash",
                    contents=quiz_prompt,
                )
                raw_text = getattr(response, "text", None) or getattr(response, "content", "")
                stripped = raw_text.strip()

                if stripped.startswith("```"):
                    stripped = re.sub(r"^```[a-zA-Z]*\n", "", stripped)
                    stripped = re.sub(r"```$", "", stripped)
                    stripped = stripped.strip()

                quiz_data = json.loads(stripped)

                # Add index from 1 to 10
                for i, item in enumerate(quiz_data):
                    item["index"] = i + 1

                topic_number = len(all_quizzes) + 1
                await db.quizzes.insert_one({
                    "course_id": obj_id,
                    "topic": topic,
                    "topic_number": topic_number,
                    "quiz": quiz_data,
                    "created_at": datetime.utcnow()
                })

                all_quizzes.append({
                    "topic_number": topic_number,
                    "quiz": quiz_data
                })

            except Exception as e:
                all_quizzes.append({
                    "topic": topic,
                    "error": f"Quiz generation failed: {e}"
                })

        return {"quizzes": all_quizzes}

    except Exception as err:
        raise HTTPException(status_code=500, detail=f"Internal error: {err}")

@app.get("/courses/{course_id}/quizzes", summary="Fetch all saved quizzes for a course")
async def get_quizzes(course_id: str):
    try:
        obj_id = ObjectId(course_id)
        quizzes_cursor = db.quizzes.find({"course_id": obj_id})
        quizzes = []
        async for quiz in quizzes_cursor:
            quiz["_id"] = str(quiz["_id"])
            quiz["course_id"] = str(quiz["course_id"])
            quizzes.append(quiz)
        return {"quizzes": quizzes}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to retrieve quizzes: {e}")


@app.post("/quizzes/{quiz_id}/attempt", summary="Submit a question-level quiz attempt")
async def submit_quiz_attempt(
    quiz_id: str,
    user_id: str = Body(...),
    topic_number: int = Body(...),
    responses: List[dict] = Body(...)
):
    try:
        obj_id = ObjectId(quiz_id)
        quiz = await db.quizzes.find_one({"_id": obj_id})
        if not quiz:
            raise HTTPException(status_code=404, detail="Quiz not found")

        enriched_responses = []
        for i, r in enumerate(responses):
            rating = r.get("user_rating")
            delay_days = rating_schedule.get(rating, 1)
            next_due_date = datetime.utcnow().date() + timedelta(days=delay_days)

            enriched_responses.append({
                "question_number": i + 1,
                "question": r.get("question"),
                "answer": r.get("answer"),
                "user_rating": rating,
                "next_due_date": str(next_due_date)
            })

        attempt_doc = {
            "quiz_id": obj_id,
            "user_id": user_id,
            "topic_number": topic_number,
            "responses": enriched_responses,
            "taken_at": datetime.utcnow()
        }

        result = await db.attempts.insert_one(attempt_doc)
        return {"status": "success", "attempt_id": str(result.inserted_id)}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save attempt: {e}")


# ────────────────────────────────────────────────────
# Fetch Upcoming Due Questions for a User
# ────────────────────────────────────────────────────

@app.get("/users/{user_id}/schedule", summary="Get upcoming questions due for review")
async def get_due_schedule(user_id: str):
    try:
        today = datetime.utcnow().date()
        upcoming = []

        cursor = db.attempts.find({"user_id": user_id})
        async for attempt in cursor:
            for resp in attempt.get("responses", []):
                due_date = datetime.strptime(resp["next_due_date"], "%Y-%m-%d").date()
                if due_date <= today:
                    upcoming.append({
                        "attempt_id": str(attempt["_id"]),
                        "quiz_id": str(attempt["quiz_id"]),
                        "topic_number": attempt["topic_number"],
                        "question_number": resp.get("question_number"),
                        "question": resp["question"],
                        "answer": resp["answer"],
                        "due_date": resp["next_due_date"]
                    })

        return {"due_questions": upcoming}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch schedule: {e}")


# ────────────────────────────────────────────────────
# PATCH: Update user rating for a specific question in an attempt
# ────────────────────────────────────────────────────

@app.patch("/attempts/{attempt_id}/update", summary="Update rating for a question")
async def update_question_rating(
    attempt_id: str,
    topic_number: int = Body(...),
    question_number: int = Body(...),
    new_rating: RatingEnum = Body(...)
):
    try:
        obj_id = ObjectId(attempt_id)
        attempt = await db.attempts.find_one({"_id": obj_id})
        if not attempt:
            raise HTTPException(status_code=404, detail="Attempt not found")

        updated = False
        for response in attempt["responses"]:
            if response.get("question_number") == question_number:
                delay_days = rating_schedule.get(new_rating, 1)
                next_due = datetime.utcnow().date() + timedelta(days=delay_days)
                response["user_rating"] = new_rating
                response["next_due_date"] = str(next_due)
                updated = True
                break

        if not updated:
            raise HTTPException(status_code=404, detail="Question not found in this attempt")

        await db.attempts.update_one({"_id": obj_id}, {"$set": {"responses": attempt["responses"]}})
        return {"status": "updated", "question_number": question_number, "new_due": str(next_due)}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to update rating: {e}")

@app.post("/courses/{course_id}/notes", summary="Upload notes for a specific topic")
async def upload_notes(
    course_id: str,
    title: str = Form(...),
    content: str = Form(""),  # Made optional
    topic_number: int = Form(...),
    image: Optional[UploadFile] = File(None),  # Made optional
    file: Optional[UploadFile] = File(None)  # Added for PDF and DOCX files
):
    try:
        # Validate the course ID
        try:
            obj_id = ObjectId(course_id)
            course = await db.courses.find_one({"_id": obj_id})
            if not course:
                raise HTTPException(status_code=404, detail="Course not found")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid course ID: {e}")
        
        # Check if we have either content, image, or file
        if not content and not image and not file:
            raise HTTPException(
                status_code=400, 
                detail="Either notes content, an image, or a document file must be provided"
            )
        
        # Create note document
        note = {
            "course_id": obj_id,
            "topic_number": topic_number,
            "title": title,
            "content": content,
            "created_at": datetime.utcnow()
        }
        
        # Handle file upload if provided (PDF, DOCX)
        extracted_text = ""
        file_id = None
        if file:
            # Read file data
            file_contents = await file.read()
            file_type = file.content_type
            
            # Validate file type
            if not file_type or not (file_type == "application/pdf" or 
                                    file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or
                                    file_type == "application/msword" or
                                    file_type == "text/plain"):
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid file type: {file_type}. Only PDF, DOCX, DOC, and TXT files are allowed."
                )
            
            # Extract text from the file based on its type
            if file_type == "application/pdf":
                try:
                    with io.BytesIO(file_contents) as pdf_file:
                        with pdfplumber.open(pdf_file) as pdf:
                            extracted_text = "\n".join([page.extract_text() or "" for page in pdf.pages])
                            extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                except Exception as e:
                    print(f"Error extracting text from PDF: {e}")
                    raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {e}")
            
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                try:
                    with io.BytesIO(file_contents) as docx_file:
                        doc = docx.Document(docx_file)
                        extracted_text = "\n".join([para.text for para in doc.paragraphs])
                        extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                except Exception as e:
                    print(f"Error extracting text from DOCX: {e}")
                    raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {e}")
            
            elif file_type == "text/plain":
                extracted_text = file_contents.decode('utf-8', errors='ignore')
                extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
            
            # Save file to GridFS for reference
            metadata = {
                "content_type": file_type,
                "filename": file.filename,
                "course_id": str(obj_id),
                "topic_number": topic_number
            }
            
            # Use the async upload_from_stream method
            file_id = await fs.upload_from_stream(
                file.filename,
                file_contents,
                metadata=metadata
            )
            
            # Add file reference to note
            note["file_id"] = str(file_id)
            
            # If extracted text is available, append it to content or use it as content
            if extracted_text:
                if content:
                    note["content"] = content + "\n\n" + extracted_text
                else:
                    note["content"] = extracted_text
                print(f"Extracted {len(extracted_text)} characters from file {file.filename}")
        
        # Handle image upload if provided
        if image:
            # Read image data
            contents = await image.read()
            
            # Get content type and validate it's an image
            content_type = image.content_type
            if not content_type or not content_type.startswith('image/'):
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid file type: {content_type}. Only images are allowed."
                )
            
            # Save to GridFS using the async fs bucket
            metadata = {
                "content_type": content_type,
                "filename": image.filename,
                "course_id": str(obj_id),
                "topic_number": topic_number
            }
            
            # Use the async upload_from_stream method
            image_id = await fs.upload_from_stream(
                image.filename,
                contents,
                metadata=metadata
            )
            
            # Add image reference to note
            note["image_id"] = str(image_id)
            
            # If no content provided but image is present, set a placeholder message
            if not content and not extracted_text:
                if content_type.startswith("image/"):
                    # Check if it's likely a handwritten note based on content type and filename
                    is_likely_handwritten = False
                    
                    if image.filename and any(kw in image.filename.lower() for kw in ["note", "handwritten", "scan", "hw"]):
                        is_likely_handwritten = True
                        
                    # Add appropriate message based on what we think it is
                    if is_likely_handwritten:
                        note["content"] = "Handwritten notes uploaded as image. AI will analyze the handwriting for quiz generation."
                        print(f"Creating quiz from handwritten notes image: {image.filename}")
                    else:
                        note["content"] = "Notes created from uploaded image."
                        print(f"Creating notes from image only: {image.filename}")
        
        # Save note to database
        result = await db.notes.insert_one(note)
        
        return {
            "status": "success",
            "notes_id": str(result.inserted_id),
            "image_id": note.get("image_id"),
            "file_id": note.get("file_id"),
            "extracted_text_length": len(extracted_text) if extracted_text else 0
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error uploading notes: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to upload notes: {e}")

@app.post("/courses/{course_id}/topics/{topic_number}/notes-quiz", summary="Generate multiple-choice quiz from notes")
async def generate_notes_quiz(course_id: str, topic_number: int):
    try:
        print(f"Generating notes quiz for course_id: {course_id}, topic: {topic_number}")
        
        try:
            obj_id = ObjectId(course_id)
            print(f"Converted to ObjectId: {obj_id}")
        except Exception as e:
            print(f"Error converting course_id to ObjectId: {e}")
            raise HTTPException(status_code=400, detail=f"Invalid course ID format: {e}")
            
        course = await db.courses.find_one({"_id": obj_id})
        if not course:
            print(f"Course not found: {course_id}")
            raise HTTPException(status_code=404, detail="Course not found")

        # Find the notes for this topic
        print(f"Searching for notes for topic_number: {topic_number}")
        notes_cursor = db.notes.find({
            "course_id": obj_id,
            "topic_number": topic_number
        }).sort("created_at", -1)  # Get most recent first
        
        notes = await notes_cursor.to_list(length=1)
        if not notes:
            print(f"No notes found for topic: {topic_number}")
            raise HTTPException(status_code=404, detail="No notes found for this topic")
            
        latest_note = notes[0]
        note_title = latest_note.get("title", f"Topic {topic_number} Notes")
        note_content = latest_note.get("content", "")
        print(f"Found notes: title='{note_title}', content_length={len(note_content)}")
        
        # Check if there's an image associated with the notes
        image_id = latest_note.get("image_id")
        image_data = None
        image_mime_type = None
        is_handwritten = False
        
        if image_id:
            print(f"Notes have image_id: {image_id}")
            # Check if it might be handwritten notes
            note_content_lower = note_content.lower()
            if "handwritten" in note_content_lower or "scan" in note_content_lower:
                is_handwritten = True
                print("Image likely contains handwritten content based on note description")
                
            # Find the image using the async GridFS
            try:
                # Convert string ID to ObjectId if needed
                if isinstance(image_id, str):
                    image_id = ObjectId(image_id)
                
                # Use the async GridFS to download the image
                print(f"Attempting to download image with id: {image_id}")
                
                # First check metadata
                image_info = await db.fs.files.find_one({"_id": image_id})
                if not image_info:
                    print(f"Image not found in GridFS: {image_id}")
                else:
                    # Get content type from metadata
                    if "metadata" in image_info and "content_type" in image_info["metadata"]:
                        image_mime_type = image_info["metadata"]["content_type"]
                        print(f"Image mime type from metadata: {image_mime_type}")
                        
                        # Check filename for more hints about handwritten content
                        if "metadata" in image_info and "filename" in image_info["metadata"]:
                            filename = image_info["metadata"]["filename"].lower()
                            if any(kw in filename for kw in ["note", "handwritten", "scan", "hw"]):
                                is_handwritten = True
                                print(f"Image likely contains handwritten content based on filename: {filename}")
                    else:
                        image_mime_type = "image/jpeg"  # Default if not specified
                        print(f"Using default image mime type: {image_mime_type}")
                    
                    # Get a download stream from GridFSBucket
                    grid_out = await fs.open_download_stream(image_id)
                    
                    # Read the content
                    chunks = []
                    async for chunk in grid_out:
                        chunks.append(chunk)
                    
                    # Combine all chunks
                    image_data = b''.join(chunks)
                    image_data_size = len(image_data) if image_data else 0
                    print(f"Retrieved image data, size: {image_data_size} bytes")
                    
                    if is_handwritten:
                        print("Will use enhanced handwritten note processing for quiz generation")
                
            except Exception as e:
                print(f"Error retrieving image: {e}")
                import traceback
                traceback.print_exc()
                # Continue without the image
                print("Continuing quiz generation without the image")
        
        # Check if we already have a quiz for this topic
        existing_quiz = await db.notes_quizzes.find_one({
            "course_id": obj_id,
            "topic_number": topic_number
        })
        
        if existing_quiz:
            print(f"Found existing quiz, will replace it. ID: {existing_quiz['_id']}")
            # Update existing quiz
            await db.notes_quizzes.delete_one({"_id": existing_quiz["_id"]})
        
        print("Calling notes_quiz_generator.create_notes_quiz_endpoint...")
        # Generate new quiz
        try:
            quiz_data = create_notes_quiz_endpoint(
                str(obj_id),  # Pass as string as expected by the function
                topic_number,
                note_title,
                note_content,
                image_data,
                image_mime_type
            )
            print(f"Quiz generation complete. Generated {len(quiz_data.get('questions', []))} questions")
            
        except Exception as e:
            print(f"Error in quiz generation: {e}")
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Quiz generation failed: {e}")
        
        # Check for errors in quiz generation
        if "error" in quiz_data:
            print(f"Quiz generation returned error: {quiz_data['error']}")
            raise HTTPException(status_code=500, detail=quiz_data["error"])
        
        # Convert the course_id back to ObjectId for MongoDB storage
        quiz_data["course_id"] = obj_id
        print(f"Converted course_id string back to ObjectId for storage: {quiz_data['course_id']}")
        
        # Add image_id if available
        if image_id:
            quiz_data["image_id"] = str(image_id)
            print(f"Added image_id to quiz: {quiz_data['image_id']}")
        
        # Set creation timestamp
        quiz_data["created_at"] = datetime.utcnow()
        
        # Store in database
        print("Saving quiz to database...")
        result = await db.notes_quizzes.insert_one(quiz_data)
        quiz_data["_id"] = str(result.inserted_id)
        print(f"Quiz saved with ID: {quiz_data['_id']}")
        
        # Convert ObjectId to string for JSON response
        quiz_data["course_id"] = str(quiz_data["course_id"])
        
        return quiz_data
    
    except HTTPException:
        raise
    except Exception as err:
        print(f"Error generating notes quiz: {err}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate quiz: {err}")

@app.get("/courses/{course_id}/topics/{topic_number}/notes-quiz", summary="Get the latest notes quiz for a topic")
async def get_notes_quiz(course_id: str, topic_number: int):
    try:
        print(f"Fetching notes quiz for course_id: {course_id}, topic: {topic_number}")
        
        try:
            obj_id = ObjectId(course_id)
            print(f"Converted to ObjectId: {obj_id}")
        except Exception as e:
            print(f"Error converting course_id to ObjectId: {e}")
            raise HTTPException(status_code=400, detail=f"Invalid course ID format: {e}")
        
        # Find the latest quiz for this topic
        print(f"Querying db.notes_quizzes for course_id: {obj_id}, topic_number: {topic_number}")
        
        # Check if collection exists
        collections = await db.list_collection_names()
        print(f"Available collections: {collections}")
        if "notes_quizzes" not in collections:
            print("notes_quizzes collection does not exist yet")
            raise HTTPException(status_code=404, detail="Notes quizzes collection not found")
            
        quiz = await db.notes_quizzes.find_one({
            "course_id": obj_id,
            "topic_number": topic_number
        })
        
        print(f"Query result: {quiz}")
        
        if not quiz:
            print("No quiz found for this topic")
            raise HTTPException(status_code=404, detail="No quiz found for this topic")
        
        # Convert ObjectIds to strings
        quiz["_id"] = str(quiz["_id"])
        quiz["course_id"] = str(quiz["course_id"])
        
        if "image_id" in quiz and quiz["image_id"]:
            print(f"Quiz has image_id: {quiz['image_id']}")
        
        return quiz
        
    except HTTPException:
        raise
    except Exception as err:
        print(f"Error retrieving notes quiz: {err}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to retrieve quiz: {err}")

@app.get("/images/{image_id}", summary="Get an image from GridFS")
async def get_image(image_id: str):
    try:
        print(f"Fetching image with ID: {image_id}")
        
        try:
            obj_id = ObjectId(image_id)
            print(f"Converted to ObjectId: {obj_id}")
        except Exception as e:
            print(f"Error converting image_id to ObjectId: {e}")
            raise HTTPException(status_code=400, detail=f"Invalid image ID format: {e}")
        
        try:
            # Check if the image exists in GridFS by getting file metadata
            file_info = await db.fs.files.find_one({"_id": obj_id})
            if not file_info:
                print(f"Image not found: {image_id}")
                raise HTTPException(status_code=404, detail="Image not found")
                
            print(f"Found image: {file_info.get('filename', 'unnamed')}")
            
            # Get content type from metadata
            content_type = "image/jpeg"  # Default
            if "metadata" in file_info and "content_type" in file_info["metadata"]:
                content_type = file_info["metadata"]["content_type"]
            print(f"Content type: {content_type}")
            
            # Download the file from GridFS using the async API
            # Get a download stream from GridFSBucket
            grid_out = await fs.open_download_stream(obj_id)
            
            # Read the content
            chunks = []
            async for chunk in grid_out:
                chunks.append(chunk)
            
            # Combine all chunks
            contents = b''.join(chunks)
            print(f"Read {len(contents)} bytes from GridFS")
            
            # Return the image with proper content type
            from fastapi.responses import Response
            return Response(content=contents, media_type=content_type)
            
        except Exception as e:
            print(f"Error retrieving image from GridFS: {e}")
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Failed to retrieve image: {e}")
        
    except HTTPException:
        raise
    except Exception as err:
        print(f"Error in image endpoint: {err}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Internal server error: {err}")

@app.get("/courses/{course_id}/all-quizzes", summary="List all quizzes for a course")
async def list_all_quizzes(course_id: str):
    try:
        print(f"Listing all quizzes for course: {course_id}")
        
        try:
            obj_id = ObjectId(course_id)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid course ID format: {e}")
            
        # First check the course exists
        course = await db.courses.find_one({"_id": obj_id})
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
            
        # Query all types of quizzes
        
        # 1. Pre-lecture quizzes from the quizzes collection
        pre_lecture_cursor = db.quizzes.find({"course_id": obj_id})
        pre_lecture_quizzes = []
        async for quiz in pre_lecture_cursor:
            quiz["_id"] = str(quiz["_id"])
            quiz["course_id"] = str(quiz["course_id"])
            quiz["quiz_type"] = "pre_lecture"
            pre_lecture_quizzes.append(quiz)
            
        # 2. Notes quizzes from the notes_quizzes collection
        notes_quizzes_cursor = db.notes_quizzes.find({"course_id": obj_id})
        notes_quizzes = []
        async for quiz in notes_quizzes_cursor:
            quiz["_id"] = str(quiz["_id"])
            quiz["course_id"] = str(quiz["course_id"])
            quiz["quiz_type"] = "notes"
            notes_quizzes.append(quiz)
            
        # Combine all quizzes
        all_quizzes = {
            "course_id": course_id,
            "pre_lecture_quizzes": pre_lecture_quizzes,
            "notes_quizzes": notes_quizzes,
            "total_quizzes": len(pre_lecture_quizzes) + len(notes_quizzes)
        }
        
        # Also retrieve information about topics from the roadmap
        if "roadmap" in course:
            topic_info = []
            for i, entry in enumerate(course["roadmap"]):
                topic_info.append({
                    "topic_number": i + 1,
                    "topic": entry.get("topic", f"Topic {i+1}")
                })
            all_quizzes["topics"] = topic_info
            
        return all_quizzes
        
    except HTTPException:
        raise
    except Exception as err:
        print(f"Error listing quizzes: {err}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to list quizzes: {err}")

@app.delete("/courses/{course_id}", summary="Delete a course and all associated data")
async def delete_course(course_id: str):
    """Delete a course and all its associated data including quizzes, notes, and images."""
    try:
        # Convert string ID to ObjectId
        try:
            course_oid = ObjectId(course_id)
        except Exception:
            raise HTTPException(status_code=400, detail="Invalid course ID format")
        
        # Check if course exists
        course = await db.courses.find_one({"_id": course_oid})
        if not course:
            raise HTTPException(status_code=404, detail="Course not found")
        
        # Delete all associated quizzes
        delete_quizzes_result = await db.quizzes.delete_many({"course_id": course_id})
        
        # Delete all associated notes
        notes_cursor = db.notes.find({"course_id": course_id})
        notes_to_delete = []
        image_ids_to_delete = []
        
        async for note in notes_cursor:
            notes_to_delete.append(note["_id"])
            # Collect image IDs to delete from GridFS
            if "image_id" in note and note["image_id"]:
                image_ids_to_delete.append(ObjectId(note["image_id"]))
        
        # Delete collected notes
        if notes_to_delete:
            await db.notes.delete_many({"_id": {"$in": notes_to_delete}})
        
        # Delete associated images from GridFS
        for image_id in image_ids_to_delete:
            try:
                await fs.delete(image_id)
            except Exception as e:
                print(f"Error deleting image {image_id}: {e}")
        
        # Delete notes quizzes
        await db.notes_quizzes.delete_many({"course_id": course_id})
        
        # Delete the course itself
        delete_result = await db.courses.delete_one({"_id": course_oid})
        
        if delete_result.deleted_count == 0:
            raise HTTPException(status_code=404, detail="Course not found or already deleted")
        
        return {"status": "success", "message": "Course and all associated data deleted successfully"}
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error deleting course: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete course: {str(e)}")

@app.post("/test/generate-quiz", summary="Test quiz generation with Gemini")
async def test_quiz_generation():
    try:
        # Sample notes for testing
        test_notes = """
        The Krebs cycle, also known as the citric acid cycle or tricarboxylic acid (TCA) cycle, 
        is a series of chemical reactions used by all aerobic organisms to release stored energy. 
        It occurs in the mitochondrial matrix and is a key metabolic pathway that connects 
        carbohydrate, fat, and protein metabolism.
        
        The cycle starts when acetyl-CoA, derived from pyruvate oxidation, joins with oxaloacetate 
        to form citrate. Through a series of reactions, citrate is transformed back to oxaloacetate, 
        releasing carbon dioxide and energy-rich compounds like NADH and FADH2. 
        These compounds enter the electron transport chain to produce ATP.
        
        The Krebs cycle produces:
        - 3 NADH molecules
        - 1 FADH2 molecule
        - 1 GTP (equivalent to ATP)
        - 2 CO2 molecules
        
        Key enzymes in the cycle include citrate synthase, aconitase, isocitrate dehydrogenase, 
        α-ketoglutarate dehydrogenase, succinyl-CoA synthetase, succinate dehydrogenase, 
        fumarase, and malate dehydrogenase.
        """
        
        print("Testing quiz generation with Gemini API")
        from notes_quiz_generator import generate_multiple_choice_quiz
        
        # Direct test of the quiz generation function
        result = generate_multiple_choice_quiz(
            notes_text=test_notes,
            note_title="The Krebs Cycle (Test)",
            image_content=None,
            image_mime_type=None
        )
        
        print(f"Quiz generation result: {len(result)} questions generated")
        
        if len(result) == 1 and "error" in result[0]:
            print(f"Quiz generation failed: {result[0]['error']}")
            return {"status": "error", "message": result[0]["error"]}
            
        return {
            "status": "success",
            "questions_count": len(result),
            "questions": result
        }
        
    except Exception as e:
        print(f"Error in test quiz generation: {e}")
        import traceback
        traceback.print_exc()
        return {"status": "error", "message": str(e)}

@app.post("/test/analyze-handwriting", summary="Test analyzing handwritten notes with Gemini Vision")
async def test_analyze_handwriting(
    image: UploadFile = File(...)
):
    try:
        # Read image data
        image_content = await image.read()
        image_mime_type = image.content_type
        
        if not image_mime_type or not image_mime_type.startswith('image/'):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file type: {image_mime_type}. Only images are allowed."
            )
        
        print(f"Analyzing handwritten image: {image.filename}, size: {len(image_content)} bytes")
        
        # Create a prompt for analyzing handwritten notes
        prompt = """
You are an expert at analyzing handwritten notes. 
Please carefully examine this image containing handwritten notes.

1. First, transcribe the text exactly as written
2. Then provide a brief summary of the main topics covered

Respond in this format:
TRANSCRIPTION:
[The exact text transcribed from the handwriting]

SUMMARY:
[A brief summary of the content's main topics]
"""
        
        try:
            # Import necessary libraries
            import io
            from PIL import Image
            
            # Convert bytes to PIL Image
            img = Image.open(io.BytesIO(image_content))
            
            # Create a Gemini model instance for vision
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            # Generate content with the image and prompt
            response = model.generate_content([prompt, img])
            
            # Get the response text
            result = response.text
            print("Analysis complete")
            
            return {
                "status": "success",
                "analysis": result,
                "image_size": len(image_content)
            }
            
        except Exception as e:
            print(f"Error analyzing handwriting: {e}")
            raise HTTPException(status_code=500, detail=f"Failed to analyze handwriting: {e}")
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in handwriting analysis endpoint: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {e}")

@app.post("/process-image/", summary="Process image and extract text plus questions")
async def process_image(file: UploadFile = File(...)):
    try:
        # Read the uploaded image file
        contents = await file.read()
        content_type = file.content_type
        
        if not content_type or not content_type.startswith('image/'):
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file type: {content_type}. Only images are allowed."
            )
        
        print(f"Processing image: {file.filename}, size: {len(contents)} bytes")
        
        # Import necessary libraries
        import io
        from PIL import Image
        
        # Open the image with PIL
        img = Image.open(io.BytesIO(contents))
        
        # Create a Gemini model instance for vision
        model = google.generativeai.GenerativeModel('gemini-2.0-flash')
        
        # Prepare the prompt for the model
        extract_prompt = "Extract the handwritten text from this image. Be thorough and capture all content."
        
        # Send the image and prompt to the model
        print("Extracting text from image...")
        extraction_response = model.generate_content([extract_prompt, img])
        
        # Get the extracted text from the response
        extracted_text = extraction_response.text
        print(f"Extracted text length: {len(extracted_text)} characters")
        
        # Generate 10 example questions based on the extracted text
        quiz_prompt = f"""
Based on the following extracted text from handwritten notes, create 10 multiple-choice questions.

EXTRACTED TEXT:
{extracted_text}

Each question should:
1. Test key concepts from the notes
2. Have 4 possible answer choices (A, B, C, D)
3. Have exactly one correct answer
4. Be challenging but fair

Format each question as a JSON object with these fields:
- id: question number (1-10)
- question: the full question text
- options: array of 4 possible answers
- correctAnswer: the correct answer (must be one of the options)

Return ONLY a valid JSON array with these 10 questions. No explanation or other text.
"""
        
        # Create a text model instance for quiz generation
        text_model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Generate quiz questions
        print("Generating quiz questions...")
        quiz_response = text_model.generate_content(quiz_prompt)
        quiz_text = quiz_response.text
        
        # Clean up any markdown code block formatting
        cleaned_quiz = quiz_text.strip()
        if cleaned_quiz.startswith("```"):
            cleaned_quiz = re.sub(r"^```[a-zA-Z]*\n", "", cleaned_quiz)
            cleaned_quiz = re.sub(r"```$", "", cleaned_quiz)
            cleaned_quiz = cleaned_quiz.strip()
        
        # Parse the JSON
        try:
            quiz_data = json.loads(cleaned_quiz)
            print(f"Generated {len(quiz_data)} quiz questions")
        except json.JSONDecodeError as e:
            print(f"Failed to parse quiz JSON: {e}")
            quiz_data = []
        
        # Return the results
        return {
            "success": True,
            "extracted_text": extracted_text,
            "quiz_questions": quiz_data
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error in process-image endpoint: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to process image: {str(e)}")
